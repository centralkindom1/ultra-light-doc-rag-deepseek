import sys
import os
import json
import re
import time
import requests
import urllib3
import traceback
import threading
from datetime import datetime
from typing import List, Dict, Optional, Union

# ==========================================
# 0. 全局配置与依赖检查
# ==========================================

# 禁用 urllib3 的安全警告 (针对 Win7/局域网自签名证书)
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

try:
    from docx import Document
    from sentence_transformers import SentenceTransformer
    from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                                 QHBoxLayout, QTextEdit, QPushButton, QLabel, 
                                 QFileDialog, QListWidget, QSplitter, QMessageBox, 
                                 QProgressBar, QGroupBox, QTableWidget, QTableWidgetItem, 
                                 QHeaderView, QAbstractItemView, QComboBox, QTabWidget,
                                 QStyleFactory)
    from PyQt5.QtCore import Qt, QThread, pyqtSignal, QObject
    from PyQt5.QtGui import QFont, QColor, QTextCursor, QTextCharFormat
except ImportError as e:
    print("【严重错误】缺少必要的库，请运行: pip install python-docx PyQt5 sentence-transformers requests")
    sys.exit(1)

# --- 工业级配置中心 ---
GLOBAL_CONFIG = {
    # 1. 聊天/推理 API (DeepSeek - 用于生成规则)
    "CHAT_API_URL": "https://aiplus.airchina.com.cn:18080/v1/chat/completions",
    "CHAT_MODEL": "DeepSeek-V3", 
    
    # 2. 向量化 API (BGE-M3 - 用于生产环境高精度检索)
    "EMBED_API_URL": "https://aiplus.airchina.com.cn:18080/v1/embeddings",
    "REMOTE_EMBED_MODEL": "bge-m3", # 局域网模型名称
    
    # 3. 通用鉴权
    "API_KEY": "sk-fXM4W0CdcKnNp3NVDfF85f2b90284b11AfDdF9F5627f627b",
    
    # 4. 本地模型配置
    "LOCAL_BGE_PATH": r"D:\Models\bge-small-zh-v1.5",
    
    # 5. 系统参数
    "APP_TITLE": "DRS Pro V3 - Hybrid RAG Factory (Local/Remote)",
    "MAX_SAMPLE_CHARS": 3000,   # AI 分析时的采样长度
    "CHUNK_SIZE_THRESHOLD": 100 # 简单的切片合并阈值(字符数)
}

# ==========================================
# 1. 基础设施层 (Logging & Signals)
# ==========================================

class LogSignal(QObject):
    """跨线程日志信号"""
    text_written = pyqtSignal(str, str) # content, color

class OutputStream(object):
    """重定向 stdout/stderr 到 GUI"""
    def __init__(self, signal_emitter, color="white"):
        self.emitter = signal_emitter
        self.color = color

    def write(self, text):
        if text.strip():
            self.emitter.text_written.emit(str(text), self.color)

    def flush(self): pass

# ==========================================
# 2. 核心工作线程 (Workers)
# ==========================================

class ModelLoaderWorker(QThread):
    """异步加载本地 BGE 模型"""
    finished_signal = pyqtSignal(object)
    log_signal = pyqtSignal(str)

    def run(self):
        self.log_signal.emit(f"⚙️ [System] 正在预加载本地模型: {GLOBAL_CONFIG['LOCAL_BGE_PATH']} (CPU模式)...")
        t_start = time.time()
        try:
            if os.path.exists(GLOBAL_CONFIG['LOCAL_BGE_PATH']):
                # device='cpu' 保证 Win7 兼容性
                model = SentenceTransformer(GLOBAL_CONFIG['LOCAL_BGE_PATH'], device='cpu')
                self.log_signal.emit(f"✅ [System] 本地模型加载完毕 (耗时 {time.time() - t_start:.2f}s)")
                self.finished_signal.emit(model)
            else:
                self.log_signal.emit(f"⚠️ [System] 路径不存在: {GLOBAL_CONFIG['LOCAL_BGE_PATH']} (仅远程模式可用)")
                self.finished_signal.emit(None)
        except Exception as e:
            self.log_signal.emit(f"❌ [System] 本地模型加载失败: {str(e)}")
            self.finished_signal.emit(None)

class AIAnalysisWorker(QThread):
    """调用 DeepSeek 分析文档结构"""
    result_signal = pyqtSignal(dict)
    log_signal = pyqtSignal(str)
    
    def __init__(self, text_sample):
        super().__init__()
        self.text_sample = text_sample

    def run(self):
        self.log_signal.emit(f">>> [AI] 正在请求 DeepSeek 分析 ({len(self.text_sample)} 字符)...")
        
        prompt = f"""你是一个 RAG 数据清洗专家。请分析以下 DOCX 文本片段。
任务：识别“噪音”（页眉、页脚、导航、广告）并提供 Python 正则表达式。

【文本片段】：
{self.text_sample}

【要求】：
仅返回 JSON 格式，不要 Markdown：
{{
  "noise_regex": ["正则表达式1", "正则表达式2"],
  "analysis_summary": "简短分析"
}}
"""
        headers = {'Content-Type': 'application/json', 'Authorization': f"Bearer {GLOBAL_CONFIG['API_KEY']}"}
        payload = {
            "model": GLOBAL_CONFIG['CHAT_MODEL'],
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.1,
            "stream": False
        }

        try:
            resp = requests.post(GLOBAL_CONFIG['CHAT_API_URL'], headers=headers, json=payload, verify=False, timeout=60)
            if resp.status_code == 200:
                content = resp.json()['choices'][0]['message']['content']
                # 尝试清洗 JSON
                json_str = re.search(r'\{.*\}', content, re.DOTALL)
                if json_str:
                    data = json.loads(json_str.group())
                    self.result_signal.emit(data)
                else:
                    self.log_signal.emit("❌ [AI] 无法解析 JSON")
                    self.result_signal.emit({})
            else:
                self.log_signal.emit(f"❌ [AI] API 错误: {resp.status_code} - {resp.text}")
                self.result_signal.emit({})
        except Exception as e:
            self.log_signal.emit(f"❌ [AI] 网络异常: {str(e)}")
            self.result_signal.emit({})

class BatchETLWorker(QThread):
    """
    核心批量处理线程：
    1. 读取 DOCX -> 2. 正则清洗 -> 3. 切片 -> 4. 向量化(本地/远程) -> 5. 保存 JSON
    """
    log_signal = pyqtSignal(str)
    progress_signal = pyqtSignal(int)
    finished_signal = pyqtSignal(str)

    def __init__(self, source_files: List[str], rules_json_path: str, output_path: str, 
                 use_remote: bool, local_model_instance=None):
        super().__init__()
        self.source_files = source_files
        self.rules_path = rules_json_path
        self.output_path = output_path
        self.use_remote = use_remote
        self.local_model = local_model_instance # 如果是本地模式，必须传入已加载的模型实例
        self.is_running = True

    def run(self):
        try:
            # 1. 加载规则
            self.log_signal.emit(f"📂 [Batch] 加载清洗规则: {self.rules_path}")
            with open(self.rules_path, 'r', encoding='utf-8') as f:
                rules_cfg = json.load(f)
            
            regexes = [re.compile(r) for r in rules_cfg.get("regex_rules", [])]
            self.log_signal.emit(f"✅ [Batch] 编译了 {len(regexes)} 条正则规则")

            knowledge_base = []
            total_files = len(self.source_files)
            
            for idx, file_path in enumerate(self.source_files):
                if not self.is_running: break
                
                filename = os.path.basename(file_path)
                # self.log_signal.emit(f"Processing: {filename}")
                
                # A. 提取与清洗
                doc = Document(file_path)
                clean_lines = []
                for p in doc.paragraphs:
                    text = p.text.strip()
                    if not text: continue
                    # 正则过滤
                    if any(reg.search(text) for reg in regexes):
                        continue
                    clean_lines.append(text)
                
                # B. 智能切片 (Smart Chunking - 简单合并短句)
                chunks = self.smart_chunking(clean_lines)
                if not chunks: continue

                # C. 向量化 (根据模式选择)
                vectors = []
                try:
                    if self.use_remote:
                        # 远程调用
                        vectors = self.get_remote_embeddings(chunks)
                    else:
                        # 本地调用
                        if self.local_model:
                            vectors = self.local_model.encode(chunks, show_progress_bar=False).tolist()
                        else:
                            raise Exception("本地模型未加载，无法执行本地向量化")
                except Exception as e:
                    self.log_signal.emit(f"❌ [Vector Error] 文件 {filename} 向量化失败: {e}")
                    continue

                # D. 组装数据
                for chunk_text, vec in zip(chunks, vectors):
                    knowledge_base.append({
                        "source": filename,
                        "text": chunk_text,
                        "vector": vec,
                        "model": GLOBAL_CONFIG['REMOTE_EMBED_MODEL'] if self.use_remote else "local-bge-small",
                        "processed_at": datetime.now().isoformat()
                    })

                self.progress_signal.emit(int((idx + 1) / total_files * 100))

            # 保存结果
            self.log_signal.emit(f"💾 [Batch] 正在保存 {len(knowledge_base)} 条数据到 JSON...")
            with open(self.output_path, 'w', encoding='utf-8') as f:
                json.dump(knowledge_base, f, ensure_ascii=False, indent=2)
            
            self.finished_signal.emit(f"任务完成！\n已处理文件: {total_files}\n生成条目: {len(knowledge_base)}\n保存位置: {self.output_path}")

        except Exception as e:
            self.log_signal.emit(f"❌ [Batch Critical] 批处理崩溃: {traceback.format_exc()}")
            self.finished_signal.emit("任务异常终止，请查看日志")

    def smart_chunking(self, lines: List[str]) -> List[str]:
        """合并过短的行，形成有意义的段落"""
        chunks = []
        current_chunk = ""
        
        for line in lines:
            if len(current_chunk) + len(line) < GLOBAL_CONFIG['CHUNK_SIZE_THRESHOLD']:
                current_chunk += line + " " # 简单拼接
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = line + " "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        return chunks

    def get_remote_embeddings(self, texts: List[str]) -> List[List[float]]:
        """工业级 API 调用：带分批、重试逻辑"""
        batch_size = 10 # 每次请求最多发送 10 条，防止 HTTP Body 过大
        all_embeddings = []
        
        headers = {"Authorization": f"Bearer {GLOBAL_CONFIG['API_KEY']}", "Content-Type": "application/json"}
        
        for i in range(0, len(texts), batch_size):
            batch = texts[i : i + batch_size]
            payload = {
                "model": GLOBAL_CONFIG['REMOTE_EMBED_MODEL'],
                "input": batch
            }
            
            # 简单的重试机制
            retry_count = 0
            while retry_count < 3:
                try:
                    resp = requests.post(GLOBAL_CONFIG['EMBED_API_URL'], json=payload, headers=headers, verify=False, timeout=30)
                    if resp.status_code == 200:
                        data = resp.json()['data']
                        # 确保按顺序
                        embeddings = [item['embedding'] for item in data]
                        all_embeddings.extend(embeddings)
                        break
                    else:
                        self.log_signal.emit(f"⚠️ [API Warning] 状态码 {resp.status_code}, 重试中...")
                        retry_count += 1
                        time.sleep(1)
                except Exception as e:
                    self.log_signal.emit(f"⚠️ [API Error] 网络波动: {e}, 重试中...")
                    retry_count += 1
                    time.sleep(1)
            
            if retry_count >= 3:
                raise Exception("远程 API 调用连续失败 3 次")
                
        return all_embeddings

# ==========================================
# 3. GUI 组件层
# ==========================================

class RuleDesignTab(QWidget):
    """Tab 1: 规则设计与预览"""
    def __init__(self, main_window):
        super().__init__()
        self.main = main_window # 持有主窗口引用以便访问 shared resources
        self.raw_data = [] 
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # 工具栏
        tool_layout = QHBoxLayout()
        btn_load = QPushButton("📂 1. 加载单个文档 (Load)")
        btn_load.clicked.connect(self.load_file)
        btn_ai = QPushButton("🤖 2. DeepSeek 分析 (Analyze)")
        btn_ai.clicked.connect(self.run_ai_analysis)
        btn_preview = QPushButton("👁️ 3. 清洗预览 (Preview)")
        btn_preview.clicked.connect(self.run_preview)
        btn_save = QPushButton("💾 4. 保存规则 (Save Rules)")
        btn_save.clicked.connect(self.save_rules)
        
        tool_layout.addWidget(btn_load)
        tool_layout.addWidget(btn_ai)
        tool_layout.addWidget(btn_preview)
        tool_layout.addWidget(btn_save)
        layout.addLayout(tool_layout)

        # 分割视图
        splitter = QSplitter(Qt.Horizontal)
        
        # 左：原始表格
        self.table_raw = QTableWidget()
        self.table_raw.setColumnCount(2)
        self.table_raw.setHorizontalHeaderLabels(["样式", "文本内容"])
        self.table_raw.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        splitter.addWidget(self.table_raw)
        
        # 中：规则列表
        rule_grp = QGroupBox("正则规则集")
        rule_layout = QVBoxLayout()
        self.list_rules = QListWidget()
        rule_layout.addWidget(self.list_rules)
        
        btn_add = QPushButton("+ 添加规则")
        btn_add.clicked.connect(self.add_rule)
        btn_del = QPushButton("- 删除规则")
        btn_del.clicked.connect(self.del_rule)
        rule_layout.addWidget(btn_add)
        rule_layout.addWidget(btn_del)
        rule_grp.setLayout(rule_layout)
        splitter.addWidget(rule_grp)
        
        # 右：预览结果
        self.text_preview = QTextEdit()
        self.text_preview.setPlaceholderText("清洗后的文本将显示在这里...")
        splitter.addWidget(self.text_preview)
        
        splitter.setSizes([400, 300, 400])
        layout.addWidget(splitter)

    def load_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择 Docx", "", "Word (*.docx)")
        if not path: return
        
        print(f"正在解析: {path}")
        doc = Document(path)
        self.raw_data = []
        self.table_raw.setRowCount(0)
        
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                self.raw_data.append({"style": p.style.name, "text": txt})
        
        self.table_raw.setRowCount(len(self.raw_data))
        for r, item in enumerate(self.raw_data):
            self.table_raw.setItem(r, 0, QTableWidgetItem(item['style']))
            self.table_raw.setItem(r, 1, QTableWidgetItem(item['text'][:100]))
        
        print(f"已加载 {len(self.raw_data)} 行数据。")

    def run_ai_analysis(self):
        if not self.raw_data: return
        # 采样前 3000 字
        sample = "\n".join([f"[{d['style']}] {d['text']}" for d in self.raw_data])[:GLOBAL_CONFIG['MAX_SAMPLE_CHARS']]
        
        self.worker = AIAnalysisWorker(sample)
        self.worker.log_signal.connect(print)
        self.worker.result_signal.connect(self.on_ai_result)
        self.worker.start()

    def on_ai_result(self, data):
        self.list_rules.clear()
        for r in data.get("noise_regex", []):
            self.list_rules.addItem(r)
        QMessageBox.information(self, "分析完成", f"DeepSeek 建议摘要：\n{data.get('analysis_summary', '无')}")

    def run_preview(self):
        rules = [self.list_rules.item(i).text() for i in range(self.list_rules.count())]
        compiled = [re.compile(r) for r in rules]
        
        clean_lines = []
        for item in self.raw_data:
            if not any(pat.search(item['text']) for pat in compiled):
                clean_lines.append(item['text'])
        
        self.text_preview.setText("\n\n".join(clean_lines))
        print(f"预览更新: 原始 {len(self.raw_data)} 行 -> 清洗后 {len(clean_lines)} 行")

    def save_rules(self):
        rules = [self.list_rules.item(i).text() for i in range(self.list_rules.count())]
        path, _ = QFileDialog.getSaveFileName(self, "保存规则", "industrial_rules.json", "JSON (*.json)")
        if path:
            with open(path, 'w', encoding='utf-8') as f:
                json.dump({"regex_rules": rules, "updated_at": datetime.now().isoformat()}, f, indent=2, ensure_ascii=False)
            print(f"规则已保存至 {path}")
            # 自动通知 Batch Tab 更新路径
            self.main.tab_batch.input_rules.setText(path)

    def add_rule(self):
        item = QListWidget().item
        self.list_rules.addItem("在此输入正则...")
        idx = self.list_rules.count() - 1
        item = self.list_rules.item(idx)
        item.setFlags(item.flags() | Qt.ItemIsEditable)
        self.list_rules.editItem(item)

    def del_rule(self):
        self.list_rules.takeItem(self.list_rules.currentRow())

class BatchProcessTab(QWidget):
    """Tab 2: 批量向量化流水线"""
    def __init__(self, main_window):
        super().__init__()
        self.main = main_window
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # 配置区
        grp_cfg = QGroupBox("流水线配置 (Pipeline Config)")
        form = QVBoxLayout()
        
        # 1. 规则文件
        h1 = QHBoxLayout()
        h1.addWidget(QLabel("规则文件:"))
        self.input_rules = QPushButton("选择 JSON 规则...")
        self.input_rules.clicked.connect(self.sel_rules)
        h1.addWidget(self.input_rules, 1)
        form.addLayout(h1)
        
        # 2. 文件夹
        h2 = QHBoxLayout()
        h2.addWidget(QLabel("文档目录:"))
        self.input_dir = QPushButton("选择 DOCX 文件夹...")
        self.input_dir.clicked.connect(self.sel_dir)
        h2.addWidget(self.input_dir, 1)
        form.addLayout(h2)

        # 3. 模型选择 (关键功能)
        h3 = QHBoxLayout()
        h3.addWidget(QLabel("核心模型:"))
        self.combo_model = QComboBox()
        self.combo_model.addItem("🚀 模式A: 纯本地 CPU (BGE-Small) - 极速/离线", "local")
        self.combo_model.addItem("🌐 模式B: 局域网 API (BGE-M3) - 高精度/长文本", "remote")
        h3.addWidget(self.combo_model, 1)
        form.addLayout(h3)

        grp_cfg.setLayout(form)
        layout.addWidget(grp_cfg)

        # 操作区
        self.btn_run = QPushButton("⚡ 开始全量处理 (Start Batch ETL)")
        self.btn_run.setFixedHeight(50)
        self.btn_run.setStyleSheet("font-size: 12pt; font-weight: bold; background-color: #0078D7; color: white;")
        self.btn_run.clicked.connect(self.start_batch)
        layout.addWidget(self.btn_run)

        self.pbar = QProgressBar()
        layout.addWidget(self.pbar)
        
        layout.addStretch()

    def sel_rules(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择规则", "", "JSON (*.json)")
        if path: self.input_rules.setText(path)

    def sel_dir(self):
        path = QFileDialog.getExistingDirectory(self, "选择包含DOCX的目录")
        if path: self.input_dir.setText(path)

    def start_batch(self):
        rules = self.input_rules.text()
        src_dir = self.input_dir.text()
        
        if not os.path.exists(rules) or not os.path.exists(src_dir):
            QMessageBox.critical(self, "错误", "请先正确选择规则文件和文档目录！")
            return

        files = [os.path.join(src_dir, f) for f in os.listdir(src_dir) if f.endswith(".docx")]
        if not files:
            QMessageBox.warning(self, "警告", "目录中没有 .docx 文件")
            return

        mode = self.combo_model.currentData()
        
        # 如果选本地模型，但模型未加载
        if mode == "local" and self.main.local_model is None:
            QMessageBox.critical(self, "错误", f"本地模型加载失败或未完成。\n请检查路径: {GLOBAL_CONFIG['LOCAL_BGE_PATH']}")
            return

        # 锁定界面
        self.btn_run.setEnabled(False)
        self.pbar.setValue(0)
        
        # 启动线程
        self.worker = BatchETLWorker(
            source_files=files,
            rules_json_path=rules,
            output_path=os.path.join(src_dir, f"knowledge_base_{mode}.json"),
            use_remote=(mode == "remote"),
            local_model_instance=self.main.local_model
        )
        self.worker.log_signal.connect(print) # 输出到主控台
        self.worker.progress_signal.connect(self.pbar.setValue)
        self.worker.finished_signal.connect(self.on_finished)
        self.worker.start()

    def on_finished(self, msg):
        self.btn_run.setEnabled(True)
        QMessageBox.information(self, "批处理完成", msg)

# ==========================================
# 4. 主窗口组装
# ==========================================

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.local_model = None # 全局持有本地模型
        self.init_ui()
        self.setup_logging()
        
        # 启动即异步加载本地模型，不阻塞界面
        self.loader = ModelLoaderWorker()
        self.loader.log_signal.connect(print)
        self.loader.finished_signal.connect(self.on_model_loaded)
        self.loader.start()

    def setup_logging(self):
        self.log_signal = LogSignal()
        self.log_signal.text_written.connect(self.append_log)
        sys.stdout = OutputStream(self.log_signal, "#00FF00") # Green
        sys.stderr = OutputStream(self.log_signal, "#FF5555") # Red

    def append_log(self, text, color):
        cur = self.console.textCursor()
        cur.movePosition(QTextCursor.End)
        fmt = QTextCharFormat()
        fmt.setForeground(QColor(color))
        cur.insertText(f"[{datetime.now().strftime('%H:%M:%S')}] {text}\n", fmt)
        self.console.setTextCursor(cur)
        self.console.ensureCursorVisible()

    def on_model_loaded(self, model):
        self.local_model = model
        status = "✅ 就绪" if model else "❌ 失败 (仅可用远程模式)"
        self.setWindowTitle(f"{GLOBAL_CONFIG['APP_TITLE']} | 本地引擎: {status}")

    def init_ui(self):
        self.setWindowTitle(GLOBAL_CONFIG['APP_TITLE'])
        self.resize(1200, 800)
        
        # 样式微调
        QApplication.setStyle(QStyleFactory.create("Fusion"))
        
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)

        # 选项卡容器
        self.tabs = QTabWidget()
        self.tab_rule = RuleDesignTab(self)
        self.tab_batch = BatchProcessTab(self)
        
        self.tabs.addTab(self.tab_rule, "🛠️ 规则实验室 (Rule Studio)")
        self.tabs.addTab(self.tab_batch, "🏭 批量生产线 (Batch Factory)")
        
        main_layout.addWidget(self.tabs, stretch=3)

        # 底部全局控制台
        console_grp = QGroupBox("系统控制台 (System Console)")
        con_layout = QVBoxLayout()
        self.console = QTextEdit()
        self.console.setReadOnly(True)
        self.console.setStyleSheet("background-color: #1e1e1e; color: #00ff00; font-family: Consolas;")
        con_layout.addWidget(self.console)
        console_grp.setLayout(con_layout)
        
        main_layout.addWidget(console_grp, stretch=1)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    font = QFont("Microsoft YaHei", 9)
    app.setFont(font)
    
    win = MainWindow()
    win.show()
    sys.exit(app.exec_())